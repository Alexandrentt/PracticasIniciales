#!/usr/bin/env python3
import os
import sys
from docx import Document

def docx_to_markdown(docx_path, output_path):
    try:
        doc = Document(docx_path)
        
        markdown_content = []
        markdown_content.append(f"# {os.path.basename(docx_path).replace('.docx', '')}\n")
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Simple conversion to markdown
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.replace('Heading ', '')
                    markdown_content.append(f"{'#' * int(level)} {text}\n")
                else:
                    markdown_content.append(f"{text}\n")
            else:
                markdown_content.append("\n")
        
        # Process tables
        for table in doc.tables:
            markdown_content.append("\n| " + " | ".join([cell.text for cell in table.rows[0].cells]) + " |\n")
            markdown_content.append("|" + "---|" * len(table.rows[0].cells) + "\n")
            for row in table.rows[1:]:
                markdown_content.append("| " + " | ".join([cell.text for cell in row.cells]) + " |\n")
            markdown_content.append("\n")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(''.join(markdown_content))
        
        print(f"✅ Converted: {docx_path} → {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ Error converting {docx_path}: {e}")
        return False

def main():
    # Files to convert
    files_to_convert = [
        ("drive-download-20260405T164349Z-1-001/5. Elaboración y preparación de documentos en el proyecto/5.3. Ética y responsabilidad social/5.3.docx", "5.3.md"),
        ("drive-download-20260405T164349Z-1-001/3 Estudio de casos como herramienta en la ingeniería/3.5/3.5.docx", "3.5.md"),
        ("drive-download-20260405T164349Z-1-001/4. Estudios en los proyectos/4.1. Estudio legal/4.1.docx", "4.1.md"),
        ("drive-download-20260405T164349Z-1-001/1. Planificación de proyectos y prácticas de la ingeniería/1.3/1.3.docx", "1.3.md"),
    ]
    
    success_count = 0
    for docx_path, md_path in files_to_convert:
        if os.path.exists(docx_path):
            if docx_to_markdown(docx_path, md_path):
                success_count += 1
        else:
            print(f"⚠️  File not found: {docx_path}")
    
    print(f"\n🎉 Successfully converted {success_count}/{len(files_to_convert)} files")

if __name__ == "__main__":
    main()
